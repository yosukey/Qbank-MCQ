"""実データの代わりになるサンプル一式を ``testdata/sample/`` に作る。

実装計画 §0 は「実データをテスト資産にする」として 2025年度の問題 docx と集計 CSV を
``testdata/`` に置くことを求めている。**その実ファイルはまだリポジトリに無い**ので、
同じ様式・同じ書式規則で組んだサンプルを生成し、CLI とパイプライン全体を通せるように
しておく。実ファイルが手に入ったら、そちらを ``testdata/`` に置いて差し替えればよい。

    python tools/make_sample_data.py [出力先]

作るもの:

- ``exam_2025.docx``          問題冊子(2 段組・MS明朝/MSゴシック・上付き下付き・均等割)
- ``item_stats_2025.csv``     集計 CSV(設計書 §10.2 の書式、31 パターン + 空白)
- ``broken_*.csv``            検証チェーンを確実に踏ませる壊れた CSV(実装計画 §6)
"""

from __future__ import annotations

import csv
import random
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))  # 隣の update_golden.py
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

from qbank_mcq.core.stats import PATTERNS  # noqa: E402
from qbank_mcq.core.typing_rules import LABELS  # noqa: E402

MINCHO = "ＭＳ 明朝"
GOTHIC = "ＭＳ ゴシック"
LATIN = "Times New Roman"

N_EXAMINEES = 139
EXAM_NAME = "口腔組織学定期試験"
EXAM_DATE = "2025-08-25"
DISC_TYPE = "D_25"

#: 選択肢セットを共有する設問群。実際の作問様式では同一セットに問い方を変えた設問が
#: 体系的に使われる(設計書 §2.4)。
SETS = [
    ["エナメル質", "象牙質", "セメント質", "歯　髄", "歯根膜"],
    ["横　紋", "死　帯", "頰　骨", "導　管", "歯　堤"],
    ["Krause 小体", "Merkel 盤", "Meissner 小体", "Ruffini 終末", "自由神経終末"],
    ["胎生 3-4 週", "胎生 6-7 週", "胎生 10 週", "出生時", "生後 6 か月"],
    ["滑膜 A 型細胞", "滑膜 B 型細胞", "破骨細胞", "骨芽細胞", "セメント芽細胞"],
]

#: ``(設問文の部品, 正答)``。設問文は ``(テキスト, 書式)`` の並びで書く。
#: 書式は "", "gothic", "bold", "italic", "sup", "sub"。
QUESTIONS = [
    # セット0
    ([("最も硬い組織はどれか。1つ選べ。", "")], "a", 0),
    ([("血管が分布するのはどれか。2つ選べ。", "")], "de", 0),
    ([("酸に溶け", ""), ("ない", "gothic"), ("のはどれか。1つ選べ。", "")], "e", 0),
    ([("上皮由来のものをすべて選べ。", "")], "a", 0),
    # セット1(均等割)
    ([("横紋が見られるのはどれか。1つ選べ。", "")], "a", 1),
    ([("歯の発生に関わるのはどれか。1つ選べ。", "")], "e", 1),
    ([("骨に含まれ", ""), ("ない", "gothic"), ("のはどれか。2つ選べ。", "")], "ad", 1),
    # セット2(ラテン文字を含む用語)
    ([("触覚受容器はどれか。3つ選べ。", "")], "abc", 2),
    ([("有髄神経終末で", ""), ("ない", "gothic"), ("のはどれか。1つ選べ。", "")], "e", 2),
    ([("圧覚に関与するものをすべて選べ。", "")], "ad", 2),
    # セット3(時期)
    ([("歯堤が形成される時期はどれか。1つ選べ。", "")], "b", 3),
    ([("エナメル質の石灰化が始まる時期はどれか。1つ選べ。", "")], "c", 3),
    # セット4(細胞)
    ([("骨吸収を担うのはどれか。1つ選べ。", "")], "c", 4),
    (
        [("ATP", "bold"), (" を必要とし", ""), ("ない", "gothic"), ("のはどれか。1つ選べ。", "")],
        "a",
        4,
    ),
    (
        [
            ("Ca", ""),
            ("2+", "sup"),
            (" と PO", ""),
            ("4", "sub"),
            ("3-", "sup"),
            (" を沈着させるのはどれか。2つ選べ。", ""),
        ],
        "de",
        4,
    ),
]


def _run(paragraph, text: str, style: str = ""):
    run = paragraph.add_run(text)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), GOTHIC if style == "gothic" else MINCHO)
    rFonts.set(qn("w:ascii"), LATIN)
    rFonts.set(qn("w:hAnsi"), LATIN)
    if style == "bold":
        run.font.bold = True
    elif style == "italic":
        run.font.italic = True
    elif style == "sup":
        run.font.superscript = True
    elif style == "sub":
        run.font.subscript = True
    return run


def write_docx(path: Path, rng: random.Random) -> list[str]:
    """問題 docx を書き、出題順の正答(印字記号)を返す。"""
    doc = Document()
    _run(doc.add_paragraph(), f"{EXAM_NAME}　{EXAM_DATE}")
    _run(doc.add_paragraph(), "学年　　番号　　氏名")

    correct_by_position: list[str] = []
    for position, (parts, correct, set_index) in enumerate(QUESTIONS, start=1):
        items = list(SETS[set_index])
        order = list(range(5))
        rng.shuffle(order)
        shuffled = [items[i] for i in order]
        # 正答の印字記号を並べ替え後の位置に写す。
        new_correct = "".join(sorted(LABELS[order.index(LABELS.index(c))] for c in correct))
        correct_by_position.append(new_correct)

        p = doc.add_paragraph()
        _run(p, f"{position}　")
        for text, style in parts:
            # 1 語を複数 run に割る癖を再現する(設計書 §5.1-2)。
            if len(text) > 6 and not style:
                half = len(text) // 2
                _run(p, text[:half], style)
                _run(p, text[half:], style)
            else:
                _run(p, text, style)

        for label, item in zip("ａｂｃｄｅ", shuffled):
            cp = doc.add_paragraph()
            _run(cp, f"{label}　{item}")

    _run(doc.add_paragraph(), f"＜以上 {len(QUESTIONS)} 設問＞")
    doc.add_paragraph("1")
    doc.save(str(path))
    return correct_by_position


def _counts_for(correct: str, rng: random.Random) -> dict[str, int]:
    """それらしい度数を作る。合計は必ず ``N_EXAMINEES``。"""
    weights: dict[str, float] = {}
    correct_set = set(correct)
    for pattern in PATTERNS:
        if len(pattern) != len(correct):
            base = 0.4  # 指示個数違反は少数
        else:
            base = 4.0
        overlap = len(correct_set & set(pattern))
        weights[pattern] = base * (1.0 + overlap * 3.0)
    weights[correct] *= 22.0
    weights[""] = 0.6  # 無回答

    keys = list(weights)
    picks = rng.choices(keys, weights=[weights[k] for k in keys], k=N_EXAMINEES)
    counts: dict[str, int] = {}
    for key in picks:
        counts[key] = counts.get(key, 0) + 1
    return counts


def write_stats_csv(path: Path, corrects: list[str], rng: random.Random) -> None:
    rows: list[list[object]] = []
    for position, correct in enumerate(corrects, start=1):
        counts = _counts_for(correct, rng)
        n_correct = counts.get(correct, 0)
        # 識別係数は 1/floor(N×0.25) 刻み(設計書 §9.2-6)。
        step = 1 / (N_EXAMINEES // 4)
        disc = round(rng.randint(2, 22) * step, 3)
        rows.append(
            [
                position,
                correct,
                round(n_correct / N_EXAMINEES, 4),
                n_correct,
                disc,
                *[counts.get(p, 0) for p in PATTERNS],
                counts.get("", 0),
            ]
        )

    with path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.writer(fh, lineterminator="\r\n")
        writer.writerow(["#試験名", EXAM_NAME])
        writer.writerow(["#試験日", EXAM_DATE])
        writer.writerow(["#受験者数", N_EXAMINEES])
        writer.writerow(["#識別係数定義", DISC_TYPE])
        writer.writerow(["問題", "正答肢", "正答率", "正答数", "識別係数", *PATTERNS, "空白"])
        writer.writerows(rows)


def write_broken_variants(source: Path, out_dir: Path) -> list[Path]:
    """検証チェーンを確実に踏ませる壊れた CSV(実装計画 §6「異常系」)。"""
    text = source.read_text(encoding="utf-8-sig")
    lines = text.splitlines()
    header_index = next(i for i, line in enumerate(lines) if line.startswith("問題,"))
    made: list[Path] = []

    def save(name: str, new_lines: list[str]) -> None:
        path = out_dir / name
        path.write_text("\r\n".join(new_lines) + "\r\n", encoding="utf-8-sig")
        made.append(path)

    # 1. 列欠落: 最後の度数列(空白)を全行から削る
    save(
        "broken_missing_column.csv",
        [line.rsplit(",", 1)[0] if i >= header_index else line for i, line in enumerate(lines)],
    )

    # 2. 正答不一致: 1 行目の正答肢を別のものに差し替える
    changed = list(lines)
    fields = changed[header_index + 1].split(",")
    fields[1] = "b" if fields[1] != "b" else "c"
    changed[header_index + 1] = ",".join(fields)
    save("broken_wrong_correct.csv", changed)

    # 3. 人数ではなく割合: 1 行目の度数をすべて N で割る
    ratio = list(lines)
    fields = ratio[header_index + 1].split(",")
    head, counts = fields[:5], fields[5:]
    ratio[header_index + 1] = ",".join([*head, *[f"{int(c) / N_EXAMINEES:.4f}" for c in counts]])
    save("broken_ratio_not_count.csv", ratio)

    # 4. 度数合計が設問ごとに違う: 1 行目の空白列を 1 つ増やす
    total = list(lines)
    fields = total[header_index + 1].split(",")
    fields[-1] = str(int(fields[-1]) + 1)
    total[header_index + 1] = ",".join(fields)
    save("broken_total_mismatch.csv", total)

    # 5. 行が 1 つ足りない
    save("broken_row_missing.csv", lines[:-1])

    return made


def write_golden(docx_path: Path) -> Path:
    """抽出結果をゴールデンファイルとして固定する(実装計画 §4 M3 受入条件)。

    実体は ``tools/update_golden.py``。書式がずれると回帰テストが誤って落ちるので、
    こちらで別に組み立てず必ず同じ実装を通す。
    """
    from update_golden import golden_path, render

    golden = golden_path(docx_path)
    golden.write_text(render(docx_path), encoding="utf-8")
    return golden


def main(out: str | None = None) -> int:
    out_dir = Path(out) if out else Path(__file__).resolve().parent.parent / "testdata" / "sample"
    out_dir.mkdir(parents=True, exist_ok=True)

    rng = random.Random(20250825)  # 生成物を再現可能にする
    docx_path = out_dir / "exam_2025.docx"
    corrects = write_docx(docx_path, rng)

    csv_path = out_dir / "item_stats_2025.csv"
    write_stats_csv(csv_path, corrects, rng)
    broken = write_broken_variants(csv_path, out_dir)
    golden = write_golden(docx_path)

    print(f"{docx_path}  ({len(QUESTIONS)} 設問)")
    print(f"{csv_path}")
    print(f"{golden}")
    for path in broken:
        print(f"{path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1] if len(sys.argv) > 1 else None))
