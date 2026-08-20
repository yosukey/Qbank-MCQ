"""docx 往復テスト(設計書 §5.3、実装計画 §2.1 スパイク①)。

実装計画 §2.1 は **ここが最大のリスク**としたうえで、

    docx → run結合 → 書式判定 → HTML断片 → docx出力 → 再取込 → HTML断片

を通し、**1回目と2回目のHTMLが一致する**ことを確認せよとしている。確認項目は:

- ``w:eastAsia`` から MS明朝 / MSゴシックが取れるか
- 上付き・下付きが往復するか
- 書式が同一の隣接 run の結合が正しく効くか
- 均等割の全角空白が正しく除去・復元されるか

2025年度の実 docx はまだリポジトリに無いため、**同じ書式規則で組んだ docx を
その場で生成して**往復させる。実ファイルが ``testdata/`` に置かれれば
``test_import_golden.py`` がそれも回帰対象にする。
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from itembank.core.text import normalize_choice, render_choice
from itembank.io.docx_read import (
    ParsedDocument,
    dump_runs,
    east_asia_font,
    parse_docx,
    summarize_formats,
)
from itembank.io.docx_write import (
    DEFAULT_WRITER_CONFIG,
    BookletItem,
    write_booklet,
)

MINCHO = DEFAULT_WRITER_CONFIG.mincho
GOTHIC = DEFAULT_WRITER_CONFIG.gothic
LATIN = DEFAULT_WRITER_CONFIG.latin


# ---------------------------------------------------------------------------
# 実データを模した docx を組む
# ---------------------------------------------------------------------------


def _run(paragraph, text: str, *, gothic=False, bold=False, italic=False, sup=False, sub=False):
    """出題様式どおりの run を 1 つ足す。"""
    run = paragraph.add_run(text)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), GOTHIC if gothic else MINCHO)
    rFonts.set(qn("w:ascii"), LATIN)
    rFonts.set(qn("w:hAnsi"), LATIN)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if sup:
        run.font.superscript = True
    if sub:
        run.font.subscript = True
    return run


@pytest.fixture
def source_docx(tmp_path: Path) -> Path:
    """取込元の docx。実際の様式の癖を意図的に混ぜてある。"""
    doc = Document()

    # ノイズ: ページ番号と記名欄(設計書 §5.1-5)
    doc.add_paragraph("1")
    p = doc.add_paragraph()
    _run(p, "学年　　番号　　氏名")

    # 問1: 否定形。強調が語中に入り、1 語が複数 run に割れている。
    p = doc.add_paragraph()
    _run(p, "1　酸に溶け")
    _run(p, "な", gothic=True)
    _run(p, "い", gothic=True)  # 同一書式の隣接 run(結合されるべき)
    _run(p, "のはどれか。1つ選べ。")
    for label, text in zip("ａｂｃｄｅ", ["エナメル質", "象牙質", "セメント質", "歯髄", "歯根膜"]):
        cp = doc.add_paragraph()
        _run(cp, f"{label}　{text}")

    # 問2: 均等割の選択肢(設計書 §7)
    p = doc.add_paragraph()
    _run(p, "2　横紋が見られるのはどれか。1つ選べ。")
    for label, text in zip("ａｂｃｄｅ", ["横　紋", "死　帯", "頰　骨", "導　管", "歯　堤"]):
        cp = doc.add_paragraph()
        _run(cp, f"{label}　{text}")

    # 問3: 学名(イタリック、run 分割あり)と化学式(上付き・下付き)
    p = doc.add_paragraph()
    _run(p, "3　")
    _run(p, "Strep", italic=True)
    _run(p, "tococcus", italic=True)
    _run(p, " ")
    _run(p, "mutans", italic=True)
    _run(p, " が産生するのはどれか。2つ選べ。")
    for label, parts in zip(
        "ａｂｃｄｅ",
        [
            [("H", {}), ("2", {"sub": True}), ("O", {})],
            [("Ca", {}), ("2+", {"sup": True})],
            [("PO", {}), ("4", {"sub": True}), ("3-", {"sup": True})],
            [("グルカン", {})],
            [("Krause 小体", {})],
        ],
    ):
        cp = doc.add_paragraph()
        _run(cp, f"{label}　")
        for text, kw in parts:
            _run(cp, text, **kw)

    # 問4: ラテン文字の強調は太字で表される(設計書 §5.3)
    p = doc.add_paragraph()
    _run(p, "4　")
    _run(p, "ATP", bold=True)
    _run(p, " を必要とし")
    _run(p, "ない", gothic=True)
    _run(p, "のはどれか。すべて選べ。")
    for label, text in zip("ａｂｃｄｅ", ["能動輸送", "受動拡散", "浸透", "濾過", "飲作用"]):
        cp = doc.add_paragraph()
        _run(cp, f"{label}　{text}")

    doc.add_paragraph("＜以上 4 設問＞")

    path = tmp_path / "source.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def parsed(source_docx: Path) -> ParsedDocument:
    return parse_docx(source_docx)


# ---------------------------------------------------------------------------
# 取込がそもそも正しいか
# ---------------------------------------------------------------------------


def test_all_questions_are_found(parsed: ParsedDocument) -> None:
    assert [q.number for q in parsed.questions] == [1, 2, 3, 4]
    assert all(len(q.choices) == 5 for q in parsed.questions)


def test_noise_lines_are_dropped(parsed: ParsedDocument) -> None:
    """ページ番号・記名欄・``＜以上 n 設問＞`` は設問にならない(設計書 §5.1-5)。"""
    assert any("学年" in s for s in parsed.skipped)
    assert any("以上" in s for s in parsed.skipped)
    assert "1" in parsed.skipped or any(s == "1" for s in parsed.skipped)


def test_adjacent_runs_of_the_same_format_are_merged(parsed: ParsedDocument) -> None:
    """怠ると ``<strong>な</strong><strong>い</strong>`` に割れる(設計書 §5.1-2)。"""
    assert parsed.questions[0].stem_html == "酸に溶け<strong>ない</strong>のはどれか。1つ選べ。"


def test_italic_species_name_is_one_fragment(parsed: ParsedDocument) -> None:
    stem = parsed.questions[2].stem_html
    assert stem.startswith("<i>Streptococcus</i> <i>mutans</i>")
    assert "<i>Strep</i>" not in stem


def test_superscript_and_subscript_survive(parsed: ParsedDocument) -> None:
    choices = parsed.questions[2].choice_htmls
    assert choices[0] == "H<sub>2</sub>O"
    assert choices[1] == "Ca<sup>2+</sup>"
    assert choices[2] == "PO<sub>4</sub><sup>3-</sup>"


def test_kintou_spaces_are_removed_on_import(parsed: ParsedDocument) -> None:
    assert parsed.questions[1].choice_htmls == ["横紋", "死帯", "頰骨", "導管", "歯堤"]


def test_latin_bold_is_normalized_to_strong(parsed: ParsedDocument) -> None:
    """強調は日本語ではゴシック、ラテン文字では太字。どちらも ``<strong>``(設計書 §5.3)。"""
    stem = parsed.questions[3].stem_html
    assert stem == "<strong>ATP</strong> を必要とし<strong>ない</strong>のはどれか。すべて選べ。"


def test_krause_is_not_squashed(parsed: ParsedDocument) -> None:
    """一律に空白を削ると壊れる例が実際の取込経路でも守られること。"""
    assert parsed.questions[2].choice_htmls[4] == "Krause 小体"


def test_item_types_are_derived(parsed: ParsedDocument) -> None:
    assert [q.item_type for q in parsed.questions] == ["A", "A", "X2", "XX"]


def test_emphasis_rule_is_checked_on_import(parsed: ParsedDocument) -> None:
    """問1・問4 は否定形で強調あり。問2・問3 は肯定形で強調なし。どれも規則どおり。"""
    assert [i.code for i in parsed.issues] == []


# ---------------------------------------------------------------------------
# 往復一致(受入条件)
# ---------------------------------------------------------------------------


def _to_booklet(parsed: ParsedDocument) -> list[BookletItem]:
    return [
        BookletItem(position=q.number, stem_html=q.stem_html, choices=q.choice_htmls)
        for q in parsed.questions
    ]


def test_html_is_identical_after_a_full_roundtrip(parsed: ParsedDocument, tmp_path: Path) -> None:
    """**受入条件**: 取込 → HTML → 冊子出力 → 再取込 で HTML が一致する。"""
    out = write_booklet(_to_booklet(parsed), tmp_path / "booklet.docx")
    again = parse_docx(out)

    assert [q.number for q in again.questions] == [q.number for q in parsed.questions]
    for before, after in zip(parsed.questions, again.questions):
        assert after.stem_html == before.stem_html
        assert after.choice_htmls == before.choice_htmls


def test_roundtrip_is_stable_on_a_second_pass(parsed: ParsedDocument, tmp_path: Path) -> None:
    """2 周目でも変わらない(片方向に少しずつずれていく不具合を捕まえる)。"""
    first = parse_docx(write_booklet(_to_booklet(parsed), tmp_path / "b1.docx"))
    second = parse_docx(write_booklet(_to_booklet(first), tmp_path / "b2.docx"))
    for a, b in zip(first.questions, second.questions):
        assert a.stem_html == b.stem_html
        assert a.choice_htmls == b.choice_htmls


def test_kintou_is_restored_in_the_printed_booklet(parsed: ParsedDocument, tmp_path: Path) -> None:
    """出力 docx の紙面には全角空白が戻っている(設計書 §7, §13.2)。"""
    out = write_booklet(_to_booklet(parsed), tmp_path / "booklet.docx", notice=None, title=None)
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any("横　紋" in t for t in texts)
    assert any("死　帯" in t for t in texts)
    # 取り込み直すとまた詰まる。
    assert parse_docx(out).questions[1].choice_htmls[0] == "横紋"


def test_render_override_is_honored_in_output(tmp_path: Path) -> None:
    item = BookletItem(
        position=1,
        stem_html="正しいのはどれか。1つ選べ。",
        choices=["横紋", "象牙質", "歯髄", "歯堤", "乳腺"],
        render_overrides=["横/紋", None, None, None, None],
    )
    out = write_booklet([item], tmp_path / "ov.docx", notice=None, title=None)
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any("横/紋" in t for t in texts)


def test_written_emphasis_uses_gothic_for_japanese_and_bold_for_latin(
    parsed: ParsedDocument, tmp_path: Path
) -> None:
    """設計書 §5.3 の「出力時の復元」列どおりに書けているか。"""
    out = write_booklet(_to_booklet(parsed), tmp_path / "booklet.docx", notice=None, title=None)
    doc = Document(str(out))
    japanese_gothic = latin_bold = False
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text == "ない":
                japanese_gothic = east_asia_font(run) == GOTHIC
            if run.text == "ATP":
                latin_bold = bool(run.font.bold)
    assert japanese_gothic, "日本語の強調が MSゴシックになっていない"
    assert latin_bold, "ラテン文字の強調が太字になっていない"


def test_base_format_produces_no_tags(parsed: ParsedDocument, tmp_path: Path) -> None:
    """基準書式(MS明朝 + Times New Roman)は既定でタグを生成しない(設計書 §5.3)。"""
    out = write_booklet(_to_booklet(parsed), tmp_path / "booklet.docx", notice=None, title=None)
    again = parse_docx(out)
    assert again.questions[1].stem_html == "横紋が見られるのはどれか。1つ選べ。"


def test_booklet_is_two_columns(parsed: ParsedDocument, tmp_path: Path) -> None:
    """設計書 §13.2: 2 段組。"""
    out = write_booklet(_to_booklet(parsed), tmp_path / "booklet.docx")
    doc = Document(str(out))
    body = doc.sections[-1]
    cols = body._sectPr.xpath("./w:cols")[0]
    assert cols.get(qn("w:num")) == "2"


def test_cover_notice_is_written(parsed: ParsedDocument, tmp_path: Path) -> None:
    out = write_booklet(
        _to_booklet(parsed),
        tmp_path / "b.docx",
        title="口腔組織学定期試験",
        notice="・持ち帰り不可",
    )
    texts = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "口腔組織学定期試験" in texts
    assert "・持ち帰り不可" in texts


# ---------------------------------------------------------------------------
# スパイク用の診断(実装計画 §12-3)
# ---------------------------------------------------------------------------


def test_dump_runs_reports_the_seven_fields(source_docx: Path) -> None:
    dumps = dump_runs(source_docx)
    assert dumps
    row = next(d for d in dumps if d.text == "ない")
    assert row.east_asia == GOTHIC
    assert row.ascii_name == LATIN
    assert row.bold in (None, False)


def test_summarize_formats_counts_by_kind(source_docx: Path) -> None:
    summary = summarize_formats(dump_runs(source_docx))
    assert summary[f"eastAsia={MINCHO}"] > 0
    assert summary[f"eastAsia={GOTHIC}"] > 0
    assert summary["italic"] == 3
    assert summary["sub"] == 2


def test_unexpected_formats_are_reported(tmp_path: Path) -> None:
    """下線などは許可タグに写せない。取込は続けつつ報告する(実装計画 §2.1)。"""
    doc = Document()
    p = doc.add_paragraph()
    _run(p, "1　正しいのはどれか。1つ選べ。")
    run = p.add_run("(補足)")
    run.font.underline = True
    for label, text in zip("ａｂｃｄｅ", ["あ", "い", "う", "え", "お"]):
        cp = doc.add_paragraph()
        _run(cp, f"{label}　{text}")
    path = tmp_path / "underline.docx"
    doc.save(str(path))

    result = parse_docx(path)
    assert any("underline" in u for u in result.unexpected_formats)
    # 下線は無視され、本文だけが残る。
    assert result.questions[0].stem_html.endswith("(補足)")


def test_images_are_extracted(tmp_path: Path) -> None:
    """インライン画像を保存し ``image_path`` に記録する(設計書 §5.1-9)。"""
    png = tmp_path / "fig.png"
    png.write_bytes(
        bytes.fromhex(
            "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c4"
            "890000000a49444154789c6300010000050001"
            "0d0a2db4"
            "0000000049454e44ae426082"
        )
    )
    doc = Document()
    p = doc.add_paragraph()
    _run(p, "1　図に示すのはどれか。1つ選べ。")
    doc.add_paragraph().add_run().add_picture(str(png))
    for label, text in zip("ａｂｃｄｅ", ["あ", "い", "う", "え", "お"]):
        cp = doc.add_paragraph()
        _run(cp, f"{label}　{text}")
    src = tmp_path / "withimage.docx"
    doc.save(str(src))

    images = tmp_path / "images"
    result = parse_docx(src, images_dir=images)
    assert result.questions[0].image_paths
    assert Path(result.questions[0].image_paths[0]).exists()


def test_normalize_and_render_are_inverse_for_kintou() -> None:
    for src in ["横　紋", "死　帯", "頰　骨", "導　管", "歯　堤", "乳　腺"]:
        assert render_choice(normalize_choice(src)) == src
