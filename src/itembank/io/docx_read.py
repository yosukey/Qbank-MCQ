"""問題 docx の取込パーサ(設計書 §5.1、実装計画 §4 M3)。

手順は設計書 §5.1 の 9 項目そのまま:

1. 段落の走査 — 行頭の ``数字 + 空白`` を設問の開始、全角 ``ａ``〜``ｅ`` で始まる段落を選択肢
2. run の結合 — **書式が同一の隣接 run を先に結合**してからタグ化する
3. 書式の判定 — ``w:eastAsia`` を直接読む(``run.font.name`` では日本語フォントが取れない)
4. 正規化 — 均等割の除去、NFKC
5. ノイズ除去 — ページ番号、``学年番号 氏名`` 欄、``＜以上 50 設問＞`` など
6. 指示文言の判定 — **タグ除去後**の文字列に対して行う
7. 強調規則チェック
8. 選択肢セットの同定(呼び出し側で ``core.bank.upsert_choice_set`` に渡す)
9. 図の抽出 — インライン画像を保存し ``image_path`` に記録

ヘッダ・フッタは ``document.paragraphs`` に現れないため、5 のノイズ除去は本文中の
体裁行だけを相手にすればよい。
"""

from __future__ import annotations

import hashlib
import logging
import re
import unicodedata
from collections.abc import Iterable, Sequence
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from ..core.text import Run, merge_runs, normalize_choice, normalize_stem, runs_to_html, strip_tags
from ..core.typing_rules import (
    LABELS,
    ValidationIssue,
    check_emphasis_rule,
    derive_item_type_detail,
)

log = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParserConfig:
    """フォントとノイズの判定規則。設定画面から変えられる(設計書 §14-10)。"""

    #: 強調とみなす日本語フォント名の部分文字列。
    gothic_markers: tuple[str, ...] = ("ゴシック", "Gothic")
    #: 基準書式の日本語フォント。タグを生成しない(設計書 §5.3)。
    mincho_markers: tuple[str, ...] = ("明朝", "Mincho")
    #: 体裁行として捨てる段落。
    noise_patterns: tuple[str, ...] = (
        r"^\s*[0-9]+\s*$",  # ページ番号だけの行
        r"^[\s　]*[<＜]\s*以上",  # ＜以上 50 設問＞
        r"学年.*番号.*氏名",  # 記名欄
        r"^[\s　]*氏\s*名[\s　]*$",
    )


DEFAULT_CONFIG = ParserConfig()

#: 設問の開始。行頭の「数字 + 空白」(設計書 §5.1-1)。全角数字も受ける。
#: 一致位置を run のオフセットにそのまま使うため、**NFKC の前**に当てる
#: (NFKC は文字数を変えうるのでオフセットがずれる)。
RE_QUESTION = re.compile(r"^[\s　]*([0-9０-９]{1,3})[.．、]?[\s　]+(?=\S)")
#: 選択肢。全角 ａ〜ｅ で始まる段落。半角も受けるが全角を正とする。
RE_CHOICE = re.compile(r"^[\s　]*([ａ-ｅa-e])[.．、]?[\s　]*")


@dataclass
class ParsedChoice:
    label: str
    html: str


@dataclass
class ParsedQuestion:
    """docx から取り出した 1 設問。DB にはまだ入れていない。"""

    number: int
    stem_html: str
    choices: list[ParsedChoice] = field(default_factory=list)
    image_paths: list[str] = field(default_factory=list)
    issues: list[ValidationIssue] = field(default_factory=list)

    @property
    def item_type(self) -> str | None:
        return derive_item_type_detail(self.stem_html).item_type

    @property
    def choice_htmls(self) -> list[str]:
        return [c.html for c in self.choices]

    def as_dict(self) -> dict[str, object]:
        """``--dry-run`` の JSON 出力用(実装計画 §4 M3 受入条件)。"""
        return {
            "number": self.number,
            "stem_html": self.stem_html,
            "stem_plain": strip_tags(self.stem_html),
            "item_type": self.item_type,
            "choices": [{"label": c.label, "html": c.html} for c in self.choices],
            "image_paths": list(self.image_paths),
            "issues": [
                {"code": i.code, "message": i.message, "blocking": i.blocking} for i in self.issues
            ],
        }


@dataclass
class ParsedDocument:
    questions: list[ParsedQuestion] = field(default_factory=list)
    #: サニタイズで落とした箇所(設計書 §3.1 は「除去箇所をログに残す」ことを求める)。
    removals: list[str] = field(default_factory=list)
    #: ノイズとして読み飛ばした行。
    skipped: list[str] = field(default_factory=list)
    #: 想定外の書式。実装計画 §2.1 のスパイクが「どれだけ混入しているか」を問う項目。
    unexpected_formats: list[str] = field(default_factory=list)

    def as_dict(self) -> dict[str, object]:
        return {
            "questions": [q.as_dict() for q in self.questions],
            "removals": self.removals,
            "skipped": self.skipped,
            "unexpected_formats": self.unexpected_formats,
        }

    @property
    def issues(self) -> list[ValidationIssue]:
        return [i for q in self.questions for i in q.issues]


# ---------------------------------------------------------------------------
# 書式の判定(設計書 §5.2)
# ---------------------------------------------------------------------------


def east_asia_font(run) -> str | None:
    """``w:rFonts/@w:eastAsia`` を直接読む。

    **``run.font.name`` では日本語フォントが取れない。** python-docx の ``font.name`` は
    ``w:rFonts/@w:ascii`` を読むため、MS明朝 / MSゴシックの判別に使えない(設計書 §5.2)。
    """
    rPr = run._element.rPr
    if rPr is None or rPr.rFonts is None:
        return None
    return rPr.rFonts.get(qn("w:eastAsia"))


def ascii_font(run) -> str | None:
    """``w:rFonts/@w:ascii``。ラテン文字側の書体。"""
    rPr = run._element.rPr
    if rPr is None or rPr.rFonts is None:
        return None
    return rPr.rFonts.get(qn("w:ascii"))


def _matches(name: str | None, markers: Sequence[str]) -> bool:
    if not name:
        return False
    folded = unicodedata.normalize("NFKC", name).casefold()
    return any(unicodedata.normalize("NFKC", m).casefold() in folded for m in markers)


def run_formats(run, config: ParserConfig = DEFAULT_CONFIG) -> frozenset[str]:
    """1 つの run の書式を許可タグ名の集合にする(設計書 §5.3 の対応表)。

    基準書式(MS明朝 + Times New Roman)は既定でタグを生成しない。
    強調は日本語では MSゴシック、ラテン文字では太字という二通りの表現になるため、
    **どちらも ``strong`` に正規化する**。
    """
    fmts: set[str] = set()
    if _matches(east_asia_font(run), config.gothic_markers) or bool(run.font.bold):
        fmts.add("strong")
    if run.font.italic:
        fmts.add("i")
    if run.font.superscript:
        fmts.add("sup")
    if run.font.subscript:
        fmts.add("sub")
    return frozenset(fmts)


def unexpected_formats(run) -> list[str]:
    """許可タグに写せない書式を洗い出す(実装計画 §2.1 の確認項目)。"""
    found: list[str] = []
    font = run.font
    if font.underline:
        found.append("underline")
    if font.strike:
        found.append("strike")
    if font.size is not None:
        found.append(f"size={font.size.pt:g}pt")
    try:
        rgb = font.color.rgb if font.color and font.color.type is not None else None
    except (AttributeError, ValueError):  # pragma: no cover - テーマ色など
        rgb = None
    if rgb is not None:
        found.append(f"color={rgb}")
    if font.highlight_color is not None:
        found.append(f"highlight={font.highlight_color}")
    if font.all_caps:
        found.append("all_caps")
    return found


# ---------------------------------------------------------------------------
# 段落 → run 列
# ---------------------------------------------------------------------------


def paragraph_runs(paragraph: Paragraph, config: ParserConfig = DEFAULT_CONFIG) -> list[Run]:
    """段落の run を ``(text, 書式集合)`` にし、**同一書式の隣接 run を結合**する。"""
    return merge_runs((r.text, run_formats(r, config)) for r in paragraph.runs)


def paragraph_unexpected(paragraph: Paragraph) -> list[str]:
    out: list[str] = []
    for r in paragraph.runs:
        for f in unexpected_formats(r):
            if r.text.strip():
                out.append(f"{f}: {r.text.strip()[:20]!r}")
    return out


def _drop_prefix(runs: Sequence[Run], count: int) -> list[Run]:
    """先頭から ``count`` 文字ぶんを落とす。書式は残りの run に保たれる。"""
    out: list[Run] = []
    remaining = count
    for text, fmt in runs:
        if remaining <= 0:
            out.append((text, fmt))
        elif len(text) <= remaining:
            remaining -= len(text)
        else:
            out.append((text[remaining:], fmt))
            remaining = 0
    return merge_runs(out)


def extract_images(paragraph: Paragraph, document, dest_dir: Path) -> list[str]:
    """段落中のインライン画像を保存し、パスの一覧を返す(設計書 §5.1-9)。"""
    saved: list[str] = []
    for blip in paragraph._p.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if not rid:
            continue
        try:
            part = document.part.related_parts[rid]
        except KeyError:  # pragma: no cover - 壊れた関係参照
            log.warning("画像の関係 %s が見つかりません", rid)
            continue
        blob = part.blob
        ext = Path(str(part.partname)).suffix or ".png"
        name = hashlib.sha256(blob).hexdigest()[:16] + ext
        dest_dir.mkdir(parents=True, exist_ok=True)
        path = dest_dir / name
        if not path.exists():
            path.write_bytes(blob)
        saved.append(str(path))
    return saved


# ---------------------------------------------------------------------------
# 本体
# ---------------------------------------------------------------------------


def is_noise(text: str, config: ParserConfig = DEFAULT_CONFIG) -> bool:
    """体裁行か(設計書 §5.1-5)。"""
    if not text.strip():
        return True
    return any(re.search(p, text) for p in config.noise_patterns)


def parse_docx(
    path: Path | str,
    *,
    config: ParserConfig = DEFAULT_CONFIG,
    images_dir: Path | None = None,
) -> ParsedDocument:
    """問題 docx を読み、設問と選択肢の HTML 断片を取り出す。"""
    document = Document(str(path))
    result = ParsedDocument()
    current: ParsedQuestion | None = None
    #: 設問文が複数段落にまたがる場合に、続きを設問文へ足すかどうか。
    in_choices = False

    for paragraph in document.paragraphs:
        # 図は単独の段落に置かれることが多く、その段落は本文が空になる。
        # ノイズ判定より先に拾わないと図を取りこぼす(設計書 §5.1-9)。
        images = extract_images(paragraph, document, images_dir) if images_dir else []

        raw = paragraph.text
        if is_noise(raw, config):
            if raw.strip():
                result.skipped.append(raw.strip())
            if images and current is not None:
                current.image_paths.extend(images)
            continue

        result.unexpected_formats.extend(paragraph_unexpected(paragraph))
        runs = paragraph_runs(paragraph, config)
        if not runs:
            continue

        plain = "".join(t for t, _ in runs)

        m_choice = RE_CHOICE.match(plain)
        m_question = RE_QUESTION.match(plain)

        if m_question and not (m_choice and in_choices):
            if current is not None:
                result.questions.append(current)
            number = int(unicodedata.normalize("NFKC", m_question.group(1)))
            body = _drop_prefix(runs, m_question.end())
            current = ParsedQuestion(number=number, stem_html=normalize_stem(runs_to_html(body)))
            in_choices = False
        elif m_choice and current is not None:
            label = unicodedata.normalize("NFKC", m_choice.group(1)).lower()
            body = _drop_prefix(runs, m_choice.end())
            current.choices.append(
                ParsedChoice(label=label, html=normalize_choice(runs_to_html(body)))
            )
            in_choices = True
        elif current is not None and not in_choices:
            # 設問文の続き(2 段落目以降)。
            current.stem_html = normalize_stem(current.stem_html + runs_to_html(runs))
        else:
            result.skipped.append(plain.strip())
            continue

        if current is not None and images:
            current.image_paths.extend(images)

    if current is not None:
        result.questions.append(current)

    for q in result.questions:
        q.issues.extend(validate_parsed_question(q))
    return result


def validate_parsed_question(q: ParsedQuestion) -> list[ValidationIssue]:
    """1 設問ぶんの構造チェック。正答は docx に無いのでここでは見ない。"""
    issues: list[ValidationIssue] = []

    if len(q.choices) != len(LABELS):
        issues.append(
            ValidationIssue(
                "choice_count",
                f"問{q.number}: 選択肢が {len(q.choices)} 個です(5 個必要)",
                context={"number": q.number},
            )
        )
    expected = list(LABELS[: len(q.choices)])
    got = [c.label for c in q.choices]
    if got != expected:
        issues.append(
            ValidationIssue(
                "choice_labels",
                f"問{q.number}: 選択肢の記号が a〜e の順ではありません: {''.join(got)}",
                context={"number": q.number},
            )
        )

    derivation = derive_item_type_detail(q.stem_html)
    if not derivation.ok:
        issues.append(
            ValidationIssue(
                "type_underivable",
                f"問{q.number}: {derivation.reason}",
                context={"number": q.number},
            )
        )

    for issue in check_emphasis_rule(q.stem_html, q.choice_htmls):
        issue.message = f"問{q.number}: {issue.message}"
        issue.context["number"] = q.number
        issues.append(issue)
    return issues


# ---------------------------------------------------------------------------
# スパイク用の診断(実装計画 §12-3)
# ---------------------------------------------------------------------------


@dataclass
class RunDump:
    """``(text, eastAsia, ascii, bold, italic, sup, sub)`` の 1 行。"""

    paragraph: int
    text: str
    east_asia: str | None
    ascii_name: str | None
    bold: bool | None
    italic: bool | None
    superscript: bool | None
    subscript: bool | None
    unexpected: list[str]

    def as_dict(self) -> dict[str, object]:
        return {
            "paragraph": self.paragraph,
            "text": self.text,
            "eastAsia": self.east_asia,
            "ascii": self.ascii_name,
            "bold": self.bold,
            "italic": self.italic,
            "sup": self.superscript,
            "sub": self.subscript,
            "unexpected": self.unexpected,
        }


def dump_runs(path: Path | str) -> list[RunDump]:
    """全 run の書式をそのままダンプする。

    実装計画 §12-3 のスパイク①: 「想定外の書式がどれだけあるかを目で確認する」ための
    出力。取込の前にこれを見て、許可タグを増やすか警告に留めるかを判断する。
    """
    document = Document(str(path))
    out: list[RunDump] = []
    for pi, paragraph in enumerate(document.paragraphs):
        for run in paragraph.runs:
            if not run.text:
                continue
            out.append(
                RunDump(
                    paragraph=pi,
                    text=run.text,
                    east_asia=east_asia_font(run),
                    ascii_name=ascii_font(run),
                    bold=run.font.bold,
                    italic=run.font.italic,
                    superscript=run.font.superscript,
                    subscript=run.font.subscript,
                    unexpected=unexpected_formats(run),
                )
            )
    return out


def summarize_formats(dumps: Iterable[RunDump]) -> dict[str, int]:
    """ダンプを書式ごとに数える。スパイクの結果メモに貼る用。"""
    counts: dict[str, int] = {}

    def bump(key: str) -> None:
        counts[key] = counts.get(key, 0) + 1

    for d in dumps:
        bump(f"eastAsia={d.east_asia}")
        bump(f"ascii={d.ascii_name}")
        for name, value in (
            ("bold", d.bold),
            ("italic", d.italic),
            ("sup", d.superscript),
            ("sub", d.subscript),
        ):
            if value:
                bump(name)
        for u in d.unexpected:
            bump(u.split("=")[0] if "=" in u else u)
    return dict(sorted(counts.items(), key=lambda kv: (-kv[1], kv[0])))
