"""問題冊子 docx の出力(設計書 §13.2、実装計画 §4 M5)。

要求は 3 つ:

- 表紙の注意書きをテンプレート化
- 2 段組
- **HTML タグから書式を復元し、均等割を再付与して印字**

書式の復元は設計書 §5.3 の対応表の逆をたどる。とくに強調は
「日本語では MSゴシック、ラテン文字では太字」という二通りの表現になるため、
1 つの ``<strong>`` を文字種で分けて書き出す。読み戻すとどちらも ``strong`` に
正規化され、隣接 run の結合で元の 1 つの断片に戻る。

**受入条件は往復一致**(設計書 §5.3、実装計画 §4 M5):
「docx 取込 → HTML 保存 → docx 出力 → 再取込」で同一の HTML が得られること。
``tests/test_docx_roundtrip.py`` がこれを自動で確かめる。
"""

from __future__ import annotations

from collections.abc import Sequence
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Pt

from ..core.text import Run, html_to_runs, render_choice
from ..core.typing_rules import LABELS

#: 印字用の全角記号。取込側の ``RE_CHOICE`` が全角 ａ〜ｅ を正とするため合わせる。
FULLWIDTH_LABELS = "ａｂｃｄｅ"

#: 設問番号・選択肢記号の後に置く区切り。全角空白 1 つ。
SEPARATOR = "　"


@dataclass(frozen=True)
class WriterConfig:
    """冊子の体裁。設定画面の「基準フォント」に対応する(設計書 §14-10)。"""

    #: 基準書式の日本語フォント(設計書 §5.3)。
    mincho: str = "ＭＳ 明朝"
    #: 強調の日本語フォント。
    gothic: str = "ＭＳ ゴシック"
    #: 基準書式のラテン文字フォント。
    latin: str = "Times New Roman"
    #: 本文の段数。
    columns: int = 2
    #: 本文のフォントサイズ(pt)。
    font_size_pt: float = 10.5


DEFAULT_WRITER_CONFIG = WriterConfig()

#: 表紙の注意書きの既定。テンプレートとして差し替えられる(設計書 §13.2)。
DEFAULT_NOTICE = (
    "・解答はすべてマークシートに記入すること。\n"
    "・指示された個数どおりに選択すること。指示と異なる個数の解答は無効とする。\n"
    "・問題冊子は持ち帰らないこと。"
)


@dataclass
class BookletItem:
    """冊子に載せる 1 設問。``core`` / DB から独立させて渡す。"""

    position: int
    stem_html: str
    #: **印字順**(a〜e)の選択肢 HTML。``choice_order`` は解決済みで渡す。
    choices: list[str]
    image_path: str | None = None
    #: 均等割の例外(``choice_set_items.render_override``)。印字順に並べる。
    render_overrides: list[str | None] = field(default_factory=list)

    def rendered_choices(self) -> list[str]:
        """均等割を再付与した印字用の文字列(設計書 §7)。

        ``render_overrides`` が短くても選択肢を落とさないよう、必ず長さをそろえてから
        突き合わせる(zip の暗黙の打ち切りで問題が欠けるのを防ぐ)。
        """
        overrides = list(self.render_overrides)
        overrides += [None] * (len(self.choices) - len(overrides))
        return [
            render_choice(html, ov)
            for html, ov in zip(self.choices, overrides[: len(self.choices)], strict=True)
        ]


def _is_japanese(ch: str) -> bool:
    """日本語(かな・漢字・全角記号)か。強調をゴシックと太字に振り分けるために使う。

    U+2E80 以降を日本語側とみなす。ラテン文字・数字・半角記号はすべて下回る。
    """
    return ord(ch) > 0x2E7F


def _split_by_script(text: str) -> list[tuple[str, bool]]:
    """``(部分文字列, 日本語か)`` に切り分ける。"""
    out: list[tuple[str, bool]] = []
    for ch in text:
        jp = _is_japanese(ch)
        if out and out[-1][1] == jp:
            out[-1] = (out[-1][0] + ch, jp)
        else:
            out.append((ch, jp))
    return out


def _set_fonts(run, *, east_asia: str, latin: str) -> None:
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)


def _add_formatted_run(paragraph, text: str, fmt: frozenset[str], config: WriterConfig) -> None:
    """1 つの ``(text, 書式)`` を docx の run として書き出す(設計書 §5.3 の逆変換)。"""
    strong = "strong" in fmt
    segments = _split_by_script(text) if strong else [(text, False)]

    for segment, is_jp in segments:
        run = paragraph.add_run(segment)
        run.font.size = Pt(config.font_size_pt)
        if strong:
            # 日本語はゴシックで、ラテン文字は太字で強調する。
            _set_fonts(run, east_asia=config.gothic, latin=config.latin)
            if not is_jp:
                run.font.bold = True
        else:
            _set_fonts(run, east_asia=config.mincho, latin=config.latin)
        if "i" in fmt:
            run.font.italic = True
        if "sup" in fmt:
            run.font.superscript = True
        if "sub" in fmt:
            run.font.subscript = True


def write_html_fragment(paragraph, html: str, config: WriterConfig = DEFAULT_WRITER_CONFIG) -> None:
    """HTML 断片を書式付きで段落に流し込む。"""
    for text, fmt in html_to_runs(html):
        _add_formatted_run(paragraph, text, fmt, config)


def _add_plain_run(paragraph, text: str, config: WriterConfig) -> None:
    _add_formatted_run(paragraph, text, frozenset(), config)


def _set_columns(section, count: int) -> None:
    """節の段組を設定する。python-docx が直接扱わないので XML を触る。"""
    cols = section._sectPr.xpath("./w:cols")
    if not cols:  # pragma: no cover - 既定テンプレートには必ずある
        return
    cols[0].set(qn("w:num"), str(count))


def write_booklet(
    items: Sequence[BookletItem],
    path: Path | str,
    *,
    title: str | None = None,
    notice: str | None = DEFAULT_NOTICE,
    config: WriterConfig = DEFAULT_WRITER_CONFIG,
    template: Path | str | None = None,
) -> Path:
    """問題冊子 docx を書き出す。

    表紙(1 段組)と本文(``config.columns`` 段組)を別の節に分ける。
    """
    document = Document(str(template)) if template else Document()

    if title or notice:
        cover = document.sections[0]
        _set_columns(cover, 1)
        if title:
            heading = document.add_paragraph()
            _add_plain_run(heading, title, config)
        if notice:
            for line in notice.splitlines():
                p = document.add_paragraph()
                _add_plain_run(p, line, config)
        body_section = document.add_section(WD_SECTION.NEW_PAGE)
    else:
        body_section = document.sections[0]
    _set_columns(body_section, config.columns)

    for item in items:
        stem = document.add_paragraph()
        _add_plain_run(stem, f"{item.position}{SEPARATOR}", config)
        write_html_fragment(stem, item.stem_html, config)

        if item.image_path and Path(item.image_path).exists():
            document.add_paragraph().add_run().add_picture(str(item.image_path))

        for index, choice_html in enumerate(item.rendered_choices()):
            p = document.add_paragraph()
            label = FULLWIDTH_LABELS[index] if index < len(FULLWIDTH_LABELS) else LABELS[index]
            _add_plain_run(p, f"{label}{SEPARATOR}", config)
            write_html_fragment(p, choice_html, config)

        document.add_paragraph()

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out))
    return out


def write_single_question_docx(
    stem_html: str,
    choices: Sequence[str],
    path: Path | str,
    *,
    position: int = 1,
    config: WriterConfig = DEFAULT_WRITER_CONFIG,
) -> Path:
    """1 設問だけの docx。往復テストと目視確認に使う。"""
    return write_booklet(
        [BookletItem(position=position, stem_html=stem_html, choices=list(choices))],
        path,
        title=None,
        notice=None,
        config=config,
    )


def runs_of_html(html: str) -> list[Run]:
    """``core.text.html_to_runs`` の再輸出(この層から使いやすくするため)。"""
    return html_to_runs(html)
